import os
from pathlib import Path
import logging

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

class DocumentParser:
    """Handles parsing of PDF and DOCX documents"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def parse_document(self, file_path: str) -> str:
        """
        Parse document and extract text content
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            str: Extracted text content
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension == '.docx':
                return self._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            self.logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file and extract text"""
        text = ""
        errors = []
        
        # Try PyMuPDF first (best for complex PDFs)
        if fitz:
            try:
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                doc.close()
                
                if text.strip():
                    self.logger.info("Successfully extracted text using PyMuPDF")
                    return text.strip()
            except Exception as e:
                error_msg = f"PyMuPDF failed: {str(e)}"
                self.logger.warning(error_msg)
                errors.append(error_msg)
        
        # Try pdfplumber second (good text extraction)
        if pdfplumber:
            try:
                text = ""  # Reset text
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                
                if text.strip():
                    self.logger.info("Successfully extracted text using pdfplumber")
                    return text.strip()
            except Exception as e:
                error_msg = f"pdfplumber failed: {str(e)}"
                self.logger.warning(error_msg)
                errors.append(error_msg)
        
        # Fallback to PyPDF2
        if PyPDF2:
            try:
                text = ""  # Reset text
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                
                if text.strip():
                    self.logger.info("Successfully extracted text using PyPDF2")
                    return text.strip()
            except Exception as e:
                error_msg = f"PyPDF2 failed: {str(e)}"
                self.logger.error(error_msg)
                errors.append(error_msg)
        
        # If all methods failed, provide detailed error message
        if not text.strip():
            error_details = " | ".join(errors) if errors else "No PDF parsing libraries available"
            raise ValueError(f"Could not extract text from PDF. This might be a scanned document, image-based PDF, or have complex formatting. Errors: {error_details}")
        
        return text.strip()
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file and extract text"""
        if not Document:
            raise ImportError("python-docx is required to parse DOCX files")
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("No text content found in DOCX file")
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error parsing DOCX: {str(e)}")
            raise
    
    def get_document_info(self, file_path: str) -> dict:
        """Get basic information about the document"""
        try:
            file_path = Path(file_path)
            return {
                "filename": file_path.name,
                "extension": file_path.suffix.lower(),
                "size_bytes": file_path.stat().st_size,
                "size_readable": self._format_file_size(file_path.stat().st_size)
            }
        except Exception as e:
            self.logger.error(f"Error getting document info: {str(e)}")
            return {}
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format"""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
